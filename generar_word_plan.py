import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_todo_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Styles
    C_NAVY = RGBColor(15, 23, 42)     # #0F172A
    C_GREEN = RGBColor(22, 163, 74)   # #16A34A
    C_GOLD = RGBColor(217, 119, 6)    # #D97706
    C_SLATE = RGBColor(51, 65, 85)    # #334155
    C_MUTED = RGBColor(100, 116, 139) # #64748B

    # Header Logo
    logo_path = 'Publica navojoa logotipo oficial.png'
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(12)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(3.2))

    # Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("PLAN DE ACCIÓN & TO-DO LIST OPERATIVO")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = C_NAVY

    # Subtitle / Slogan
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run("Publica Navojoa • \"Privado, personalizado y directo.\"")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(12)
    r_sub.font.bold = True
    r_sub.font.color.rgb = C_GREEN

    # Meta Info Card / Table
    table_meta = doc.add_table(rows=2, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False

    meta_data = [
        [("Objetivo del Periodo:", "Lanzamiento y activación del sistema de captación en FB + Bot de WhatsApp."),
         ("Socios / Responsables:", "Enrique Valenzuela & Mónica Obregón")],
        [("Horizonte Temporal:", "2 Semanas (Fase de Preparación y Validación Piloto)"),
         ("Meta Cuantitativa:", "Captar 150-250 contactos autorizados con Opt-in y 1er cliente piloto.")]
    ]

    for row_idx, row in enumerate(table_meta.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(3.3)
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            label, val = meta_data[row_idx][col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r_l = p.add_run(f"{label} ")
            r_l.font.bold = True
            r_l.font.size = Pt(9.5)
            r_l.font.color.rgb = C_NAVY
            r_v = p.add_run(val)
            r_v.font.size = Pt(9.5)
            r_v.font.color.rgb = C_SLATE

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # Helper to add Section Headers
    # -------------------------------------------------------------
    def add_section_header(title, color_rgb):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = color_rgb

    # -------------------------------------------------------------
    # Helper to add Subheaders
    # -------------------------------------------------------------
    def add_subheader(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = C_NAVY

    # -------------------------------------------------------------
    # Helper to add Tasks with Checkboxes
    # -------------------------------------------------------------
    def add_task(title, desc, resp="", priority="Alta"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)

        r_box = p.add_run("☐  ")
        r_box.font.bold = True
        r_box.font.size = Pt(11)
        r_box.font.color.rgb = C_GREEN

        r_title = p.add_run(f"{title}")
        r_title.font.bold = True
        r_title.font.size = Pt(10)
        r_title.font.color.rgb = C_NAVY

        if resp:
            r_resp = p.add_run(f"  [{resp} • {priority}]")
            r_resp.font.size = Pt(8.5)
            r_resp.font.bold = True
            r_resp.font.color.rgb = C_GOLD

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.45)
        p_desc.paragraph_format.space_after = Pt(6)
        r_d = p_desc.add_run(desc)
        r_d.font.size = Pt(9.5)
        r_d.font.color.rgb = C_SLATE

    # =============================================================
    # SEMANA 1
    # =============================================================
    add_section_header("📅 SEMANA 1: Activación de Canales, Activos & Captación Inicial", C_NAVY)

    # Bloque 1
    add_subheader("Fase 1.1: Configuración del Grupo de Facebook y Branding (Días 1 - 2)")
    add_task(
        "Configurar las 3 Preguntas de Admisión en Facebook",
        "Ingresar a: Administrar Grupo ➔ Preguntas a miembros potenciales. Activar las 3 preguntas: 1) Reglas, 2) Colonia en Navojoa, 3) Autorización Opt-in de WhatsApp (10 dígitos). Esto iniciará la captación en automático de cada nuevo miembro.",
        resp="Enrique / Mónica",
        priority="Urgente"
    )
    add_task(
        "Actualizar la Portada Oficial del Grupo de Facebook",
        "Subir la imagen de portada oficial (1640x856 px) con el logo oficial de Publica Navojoa, el slogan 'Privado, personalizado y directo' y una llamada a la acción clara para unirse al Club VIP.",
        resp="Enrique",
        priority="Alta"
    )
    add_task(
        "Adquirir y Activar la Línea / SIM Dedicada de WhatsApp",
        "Tener listo el chip/número telefónico exclusivo para Publica Navojoa. Descargar WhatsApp Business, configurar el perfil de empresa con el logo, descripción y horario de atención.",
        resp="Mónica",
        priority="Urgente"
    )

    # Bloque 2
    add_subheader("Fase 1.2: Despliegue de la Landing Web & Publicación Fijada (Días 3 - 4)")
    add_task(
        "Publicar la Landing Page de Registro (index.html)",
        "Alojar index.html en hosting gratuito/rápido (GitHub Pages, Netlify o publicanavojoa.com) para tener el enlace oficial listo para compartir.",
        resp="Enrique",
        priority="Alta"
    )
    add_task(
        "Publicar el Primer Anuncio Fijado en el Grupo de Facebook",
        "Redactar y fijar en la parte superior del grupo: '👑 ¡Únete al Club VIP de Ofertas de Navojoa! Recibe remates y promociones en tu celular antes que nadie → [Enlace a index.html]'.",
        resp="Enrique / Mónica",
        priority="Alta"
    )
    add_task(
        "Configurar Mensaje de Bienvenida en WhatsApp Business",
        "Configurar la respuesta automática inicial: '¡Hola [Nombre]! 👋 Bienvenido a Publica Navojoa. Ya estás en nuestra lista VIP para recibir ofertas y avisos exclusivos de la ciudad.'",
        resp="Mónica",
        priority="Media"
    )

    # Bloque 3
    add_subheader("Fase 1.3: Primera Sincronización y Validación de Datos (Días 5 - 7)")
    add_task(
        "Captura de Primeros Teléfonos al Panel Admin (admin.html)",
        "Revisar las solicitudes aprobadas en el grupo de Facebook y vaciar los primeros teléfonos capturados en el panel admin.html para mantener el registro centralizado.",
        resp="Mónica",
        priority="Alta"
    )
    add_task(
        "Prueba Piloto de Bienvenida 1 a 1 a 30 Contactos",
        "Enviar mensaje de confirmación a los primeros 30 registros autorizados para verificar que los números sean válidos y confirmar que guardaron nuestro contacto.",
        resp="Mónica",
        priority="Alta"
    )
    add_task(
        "Cierre de Semana 1 y Conteo de Registros",
        "Junta de 30 minutos entre socios para verificar el total de leads captados (meta: 50-100) y revisar ajustes para la semana 2.",
        resp="Enrique & Mónica",
        priority="Media"
    )

    # =============================================================
    # SEMANA 2
    # =============================================================
    add_section_header("📅 SEMANA 2: Primer Catálogo Piloto, Pruebas y Prospección Comercial", C_GREEN)

    # Bloque 4
    add_subheader("Fase 2.1: Creación del Primer Catálogo de Ofertas (Días 8 - 9)")
    add_task(
        "Estructurar el 'Catálogo Semanal #1 de Remates y Ofertas'",
        "Seleccionar de 4 a 6 productos atractivos (remates de muebles/decoración del archivo Excel o de comercios aliados) con fotos limpias, precio con descuento y botón de compra.",
        resp="Mónica",
        priority="Alta"
    )
    add_task(
        "Crear Lista de 15 Comercios Prospectos en Navojoa",
        "Listar 15 negocios locales con presencia activa (restaurantes, boutiques, mueblerías, estéticas) con nombre del dueño/encargado y teléfono de contacto para ofrecer el servicio.",
        resp="Enrique / Mónica",
        priority="Alta"
    )

    # Bloque 5
    add_subheader("Fase 2.2: Lanzamiento del 1er Broadcast y Prospección (Días 10 - 11)")
    add_task(
        "Ejecutar el 1er Envío Masivo por WhatsApp (Broadcast Piloto)",
        "Enviar el Catálogo Semanal a toda la base registrada. Medir cuántas personas abren el mensaje, cuántas preguntan y asegurar que cero usuarios se quejen.",
        resp="Mónica / Enrique",
        priority="Urgente"
    )
    add_task(
        "Iniciar Prospección con el Kit de Medios (media_kit.html)",
        "Enviar mensaje personalizado a los primeros 5 prospectos comerciales utilizando el guion de prospección y compartiendo el enlace al Kit de Medios con los 3 paquetes ($600, $1,200, $2,500).",
        resp="Enrique",
        priority="Alta"
    )
    add_task(
        "Publicar Dinámica Semanal en el Grupo de Facebook",
        "Post interactivo: '¿Buscas algo en específico esta semana en Navojoa? Déjanos tu comentario o entra al Club VIP para recibir las mejores promociones de comercios verificados.'",
        resp="Mónica",
        priority="Media"
    )

    # Bloque 6
    add_subheader("Fase 2.3: Cierre de Primer Cliente y Evaluación de Resultados (Días 12 - 14)")
    add_task(
        "Cierre y Ejecución del Primer Anunciante Pagado",
        "Concretar la venta del primer Paquete Bronce ($600) o Plata ($1,200). Fijar su publicación en el grupo de FB y agendar su mención en el siguiente broadcast de WhatsApp.",
        resp="Enrique & Mónica",
        priority="Urgente"
    )
    add_task(
        "Exportar Base de Datos para Meta Ads (Prueba Técnica)",
        "En admin.html, exportar el archivo CSV y cargarlo en Meta Ads Manager ➔ Públicos Personalizados para validar la creación del público similar (Lookalike) en Navojoa.",
        resp="Enrique",
        priority="Media"
    )
    add_task(
        "Revisión Mensual de KPIs y Plan de Escalamiento",
        "Evaluar resultados finales de las 2 semanas: Total de contactos captados, tasa de conversión del 1er broadcast, ingresos generados y metas para la Semana 3 en adelante.",
        resp="Enrique & Mónica",
        priority="Alta"
    )

    # -------------------------------------------------------------
    # Tabla Resumen de Roles y Responsabilidades
    # -------------------------------------------------------------
    add_section_header("👥 Matriz de Responsabilidades por Rol", C_NAVY)

    table_roles = doc.add_table(rows=4, cols=3)
    table_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_roles.autofit = False

    headers = ["Socio / Rol", "Áreas de Enfoque Principal", "Entregables Clave (2 Semanas)"]
    for idx, cell in enumerate(table_roles.rows[0].cells):
        cell.width = [Inches(1.8), Inches(2.6), Inches(2.6)][idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        r = p.add_run(headers[idx])
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    roles_data = [
        ("Enrique Valenzuela\n(Dirección Técnica & Comercial)", 
         "• Despliegue de herramientas web\n• Portada oficial de Facebook\n• Prospección a comercios VIP\n• Configuración técnica de Meta Ads", 
         "• Landing page y admin activos\n• Portada de FB publicada\n• Primeros clientes comerciales contactados\n• Exportación CSV validada"),
        ("Mónica Obregón\n(Operaciones & Comunidad)", 
         "• Moderación y admisión de FB\n• Gestión del WhatsApp Business\n• Armado del Catálogo de Ofertas\n• Captura y vaciado de leads a la base", 
         "• 3 preguntas de FB configuradas\n• Teléfonos capturados al día\n• 1er Broadcast de ofertas enviado\n• Catálogo semanal diseñado"),
        ("Trabajo Conjunto\n(Ambos Socios)", 
         "• Cierre y cobranza de clientes\n• Juntas semanales de revisión\n• Estrategia de crecimiento", 
         "• 150-250 leads autorizados\n• 1er anunciante pagado cerrado\n• Ajuste del tarifario y paquete")
    ]

    for row_idx, (r_name, r_focus, r_deliver) in enumerate(roles_data, start=1):
        row_cells = table_roles.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_text in enumerate([r_name, r_focus, r_deliver]):
            cell = row_cells[c_idx]
            cell.width = [Inches(1.8), Inches(2.6), Inches(2.6)][c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(cell_text)
            r.font.size = Pt(9)
            r.font.color.rgb = C_NAVY if c_idx == 0 else C_SLATE
            if c_idx == 0:
                r.font.bold = True

    # Callout Box: Reglas de Oro
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    tbl_callout = doc.add_table(rows=1, cols=1)
    tbl_callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_box = tbl_callout.rows[0].cells[0]
    c_box.width = Inches(6.8)
    set_cell_background(c_box, "FEF3C7") # Amber tint
    set_cell_margins(c_box, top=140, bottom=140, left=180, right=180)
    p_box = c_box.paragraphs[0]
    r_b_title = p_box.add_run("⚠️ REGLAS DE ORO OPERATIVAS PARA ESTAS 2 SEMANAS:\n")
    r_b_title.font.bold = True
    r_b_title.font.size = Pt(10)
    r_b_title.font.color.rgb = C_GOLD

    rules = (
        "1. Consentimiento 100% Obligatorio (Opt-in): Jamás enviar mensajes a personas que no hayan dejado su WhatsApp voluntariamente.\n"
        "2. No saturar: Máximo 1 o 2 mensajes por semana para evitar que los usuarios envíen 'BAJA'.\n"
        "3. El Grupo sigue libre: Publicar sigue siendo gratis para todos los miembros normales. Los paquetes pagados son para fijar publicaciones y envíos masivos."
    )
    r_b_txt = p_box.add_run(rules)
    r_b_txt.font.size = Pt(9)
    r_b_txt.font.color.rgb = C_NAVY

    output_filename = "TO_DO_LIST_PROXIMAS_2_SEMANAS_PUBLICA_NAVOJOA.docx"
    doc.save(output_filename)
    print(f"SUCCESS: Created {output_filename}")

if __name__ == "__main__":
    create_todo_docx()
